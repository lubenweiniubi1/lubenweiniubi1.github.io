"""
Insert v3 resume content into Linfeng Pan.docx template.
Replaces: Responsibilities & Achievements + Project Experience (English + Chinese)
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

doc = Document('/Users/panlinfeng/workspace/course-1-basic-stuff/简历/Linfeng Pan.docx')

FONT = 'Arial Body'
SIZE = Pt(9)

def make_para(doc, text, bold=False, style='Normal', font=FONT, size=SIZE):
    """Create a paragraph with consistent formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = size
    run.bold = bold
    return p

def make_list_item(doc, text, style='List Paragraph1'):
    """Create a bullet list item."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = SIZE
    return p

def make_heading1(doc, text):
    """Create a Heading 1 paragraph (10pt)."""
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def insert_after(doc, ref_para, new_para):
    """Insert new_para's XML element after ref_para's XML element."""
    ref_para._element.addnext(new_para._element)

def delete_range(doc, start_idx, end_idx):
    """Delete paragraphs from start_idx to end_idx (inclusive)."""
    body = doc.element.body
    for i in range(end_idx, start_idx - 1, -1):
        p = doc.paragraphs[i]
        body.remove(p._element)

def find_heading(doc, text, start=0):
    """Find paragraph index by exact Heading text."""
    for i in range(start, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == text:
            return i
    return -1

def find_next_heading1(doc, start):
    """Find next Heading 1 after start index."""
    for i in range(start + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name == 'Heading 1':
            return i
    return len(doc.paragraphs)

# ============================================================
# STEP 1: English "Responsibilities & Achievements"
# ============================================================
resp_start = find_heading(doc, 'Responsibilities & Achievements')
proj_start = find_heading(doc, 'Project Experience')
print(f"English: Responsibilities at P{resp_start}, Project Experience at P{proj_start}")

# Delete old content between Responsibilities heading and Project Experience heading
# (the paragraphs between them)
old_resp_end = proj_start - 1
delete_range(doc, resp_start + 1, old_resp_end)

# Now insert new content right after the Responsibilities heading

def insert_paras_after(after_para, para_list):
    """Insert a list of paragraphs after after_para (in order).
    para_list is a list of paragraph objects already added to the document."""
    # Insert in reverse order so they appear in correct order
    ref_elem = after_para._element
    for p in reversed(para_list):
        ref_elem.addnext(p._element)

# ---- Build English Responsibilities content ----
resp_paras = []

# Thoughtworks header
resp_paras.append(make_para(doc, 'Thoughtworks | Frontend Engineer | 2021 – 2025', bold=True))
resp_paras.append(make_para(doc, ''))

# REA Group
resp_paras.append(make_para(doc, 'REA Group — Full-Stack Development & AI Workflow (React + Node.js/Express)', bold=True))
resp_paras.append(make_list_item(doc, 'Developed frontend features and maintained Node.js/Express backend services for Australia\'s largest real estate platform'))
resp_paras.append(make_list_item(doc, 'Designed and implemented custom Agent Skills in Codex to standardize the development pipeline:'))
resp_paras.append(make_list_item(doc, '    • Spec Skill — decompose requirements into structured, executable spec documents'))
resp_paras.append(make_list_item(doc, '    • Unit Test Skill — auto-generate test cases for new features and bug fixes'))
resp_paras.append(make_list_item(doc, '    • Issue Skill — publish structured issue reports documenting root cause, fix approach, and prevention measures'))
resp_paras.append(make_list_item(doc, '    • Code Review Skill — automate pre-review checks before peer review'))
resp_paras.append(make_list_item(doc, 'Owned the full delivery lifecycle: spec → development → unit tests → issue publishing → code review → self-service production release (via pre-configured CI/CD pipeline)'))
resp_paras.append(make_list_item(doc, 'Maintained cross-team code health monitoring, driving code standard adoption and continuous quality improvement'))
resp_paras.append(make_para(doc, ''))

# Porsche
resp_paras.append(make_para(doc, 'Porsche Dealership System — Multi-Platform Development (React + TypeScript + Taro)', bold=True))
resp_paras.append(make_list_item(doc, 'Developed core features across 4 platforms: WeChat Mini-Program, Douyin Mini-Program, Web Admin, and APP H5 hybrid'))
resp_paras.append(make_list_item(doc, 'Built the vehicle intelligent recommendation module (backend-driven), new car catalog, and vehicle series modules for the Mini-Program'))
resp_paras.append(make_list_item(doc, 'Developed the Admin-side dealership management module covering dealer listing, audit workflows, permission allocation, and data dashboards; integrated Tencent Cloud Video for dealer training content'))
resp_paras.append(make_list_item(doc, 'Led the WeChat → Douyin Mini-Program migration: resolved Taro API incompatibilities and re-implemented animation modules for Douyin\'s runtime, achieving ~70% code reuse'))
resp_paras.append(make_list_item(doc, 'Contributed to a cross-platform component library (Taro + Lerna/Yarn Workspace Monorepo), building 15+ reusable business components shared across all 4 platforms'))
resp_paras.append(make_list_item(doc, 'Enforced ESLint standards and Code Review practices, reducing team-wide production defects by an estimated 50%'))
resp_paras.append(make_para(doc, ''))

# Leasing
resp_paras.append(make_para(doc, 'Vehicle Leasing System — Mini-Program + Web (React + Jest + React Testing Library)', bold=True))
resp_paras.append(make_list_item(doc, 'Built the Web order management interface (order listing, filtering, detail views) and the Mini-Program vehicle selection flow'))
resp_paras.append(make_list_item(doc, 'Wrote unit and integration tests (Jest + React Testing Library) for core leasing workflows'))
resp_paras.append(make_para(doc, ''))

# Internal Tools
resp_paras.append(make_para(doc, 'Internal Tools — KPI Dashboard / Contract Management (React + ECharts)', bold=True))
resp_paras.append(make_list_item(doc, 'Developed a KPI visualization dashboard with ECharts, replacing manual Excel reporting with real-time business metric charts'))
resp_paras.append(make_list_item(doc, 'Built the contract management system: contract listing, multi-condition filtering, and batch operations'))
resp_paras.append(make_list_item(doc, 'Refactored legacy frontend codebases, reducing bundle size by an estimated 35% and improving maintainability'))
resp_paras.append(make_para(doc, ''))

# Feifeng
resp_paras.append(make_para(doc, 'Xi\'an Feifeng Technology | Frontend Engineer | 2020 – 2021', bold=True))
resp_paras.append(make_list_item(doc, 'Developed a rural sanitation management Mini-Program (native WeChat) for China\'s rural revitalization initiative: integrated IoT toilet sensors for waste level monitoring, implemented dynamic route planning for collection vehicles, and built real-time GPS tracking with navigation'))
resp_paras.append(make_list_item(doc, 'Independently delivered a full-stack internal management system (React + Spring Boot + MongoDB) — from database modeling to frontend UI'))

# Insert after Responsibilities heading
insert_paras_after(doc.paragraphs[resp_start], resp_paras)
print(f"Inserted {len(resp_paras)} paragraphs for English Responsibilities")

# ============================================================
# STEP 2: English "Project Experience"
# ============================================================
# Rebuild indices since we modified the document
# The Project Experience heading has shifted
proj_idx = find_heading(doc, 'Project Experience')
print(f"Project Experience now at P{proj_idx}")

# Find where Chinese section starts ("个人背景" heading)
cn_start = find_heading(doc, '个人背景')
print(f"Chinese section starts at P{cn_start}")

# Delete old project content (between Project Experience heading and Chinese section)
delete_range(doc, proj_idx + 1, cn_start - 1)

# Re-find indices
proj_idx = find_heading(doc, 'Project Experience')

# ---- Build English Project Experience content ----
proj_paras = []

# === Project 1: Porsche ===
proj_paras.append(make_para(doc, '2022 – 2025       Porsche Multi-Platform Dealership System', bold=True))
proj_paras.append(make_para(doc, 'Project Description: Developed core business modules for Porsche\'s nationwide dealership network across 4 platforms (WeChat Mini-Program, Douyin Mini-Program, Web Admin, H5). The system covers vehicle recommendation, new car catalog, dealership management, and dealer training.', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Main responsibilities:', bold=True))
proj_paras.append(make_list_item(doc, 'Architected and developed core business modules across 4 platforms, ensuring consistent UX and shared business logic'))
proj_paras.append(make_list_item(doc, 'Built the vehicle intelligent recommendation flow (backend API-driven, frontend display + interaction)'))
proj_paras.append(make_list_item(doc, 'Developed Admin-side dealership management with full CRUD, audit pipeline, and role-based permission control'))
proj_paras.append(make_list_item(doc, 'Integrated Tencent Cloud Video for dealer-facing training content delivery'))
proj_paras.append(make_list_item(doc, 'Led WeChat → Douyin Mini-Program migration: resolved Taro API incompatibilities and re-implemented animation modules, achieving ~70% code reuse'))
proj_paras.append(make_list_item(doc, 'Co-built cross-platform component library (Taro + Lerna Monorepo) with 15+ shared business components, eliminating redundant per-platform UI development'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Project Outcome：', bold=True))
proj_paras.append(make_para(doc, 'Multi-platform delivery cycle reduced by an estimated 40% through component reuse. UI consistency reached 100% across all 4 platforms.'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Core Technologies：React, TypeScript, Redux Toolkit, TaroJS, WeChat Mini-Program, Douyin Mini-Program, H5, Tencent Cloud Video', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, ''))

# === Project 2: REA Group ===
proj_paras.append(make_para(doc, '2023 – 2025       REA Group — AI-Driven Full-Stack Development', bold=True))
proj_paras.append(make_para(doc, 'Project Description: Maintained and developed features for Australia\'s largest real estate platform (realcommercial.com.au / realestate.com.au), serving millions of users. Built a custom AI-assisted development pipeline using self-defined Agent Skills, covering the full delivery cycle from spec decomposition to production release.', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Main responsibilities:', bold=True))
proj_paras.append(make_list_item(doc, 'Developed frontend features (React + TypeScript) and maintained backend Node.js/Express services'))
proj_paras.append(make_list_item(doc, 'Designed 4 custom Agent Skills: Spec (requirement decomposition), Unit Test (auto-generate Jest tests), Issue (structured root-cause reports), Code Review (automated pre-review checks)'))
proj_paras.append(make_list_item(doc, 'Owned the full delivery pipeline: Spec → Development → Unit Tests → Issue Publishing → Code Review → Self-Service Release (via pre-configured CI/CD)'))
proj_paras.append(make_list_item(doc, 'All collaboration conducted in English across Australian, Indian, and Chinese teams'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Project Outcome：', bold=True))
proj_paras.append(make_para(doc, 'Transformed a manual, fragmented development process into a reproducible, quality-gated AI-assisted delivery system. Every change had mandatory test coverage, documented rationale, and formal review before reaching production.'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Core Technologies：React, TypeScript, Node.js, Express, Jest, Codex (AI-assisted)', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, ''))

# === Project 3: Rural IoT ===
proj_paras.append(make_para(doc, '2020 – 2021       Rural Sanitation IoT Management Mini-Program', bold=True))
proj_paras.append(make_para(doc, 'Project Description: Part of China\'s rural revitalization initiative — IoT sensors installed in rural public toilets monitored waste levels in real time. The Mini-Program enabled government workers to view which facilities needed servicing and follow dynamically planned collection routes.', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Main responsibilities:', bold=True))
proj_paras.append(make_list_item(doc, 'Developed the entire Mini-Program natively using WeChat\'s framework'))
proj_paras.append(make_list_item(doc, 'Integrated IoT sensor data to display real-time waste level monitoring across facilities'))
proj_paras.append(make_list_item(doc, 'Implemented a dynamic route planning algorithm to generate optimal collection paths based on sensor priority and geographic proximity'))
proj_paras.append(make_list_item(doc, 'Built real-time GPS tracking and turn-by-turn navigation for collection vehicle drivers'))
proj_paras.append(make_list_item(doc, 'Delivered the full system independently from requirements to deployment'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Project Outcome：', bold=True))
proj_paras.append(make_para(doc, 'Enabled data-driven sanitation management, replacing previously manual/schedule-based collection with sensor-driven dynamic routing.'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Core Technologies：Native WeChat Mini-Program, IoT sensor integration, GPS tracking, Route Planning', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, ''))

# === Project 4: Full-Stack Internal System ===
proj_paras.append(make_para(doc, '2020 – 2021       Full-Stack Internal Management System', bold=True))
proj_paras.append(make_para(doc, 'Project Description: Independently architected and delivered a complete internal management system covering the full stack from database design to frontend UI.', bold=True))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Main responsibilities:', bold=True))
proj_paras.append(make_list_item(doc, 'Designed the MongoDB data model and built RESTful APIs with Spring Boot'))
proj_paras.append(make_list_item(doc, 'Developed the React frontend with full CRUD functionality'))
proj_paras.append(make_list_item(doc, 'Handled deployment and production onboarding'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Project Outcome：', bold=True))
proj_paras.append(make_para(doc, 'Delivered the complete full-stack system independently, from zero to production.'))
proj_paras.append(make_para(doc, ''))
proj_paras.append(make_para(doc, 'Core Technologies：React, Spring Boot, MongoDB', bold=True))

# Insert after Project Experience heading
insert_paras_after(doc.paragraphs[proj_idx], proj_paras)
print(f"Inserted {len(proj_paras)} paragraphs for English Project Experience")

# ============================================================
# STEP 3: Chinese "职位及成就"
# ============================================================
cn_resp_start = find_heading(doc, '职位及成就')
cn_proj_start = find_heading(doc, '项目经验')
print(f"Chinese: 职位及成就 at P{cn_resp_start}, 项目经验 at P{cn_proj_start}")

# Delete old content
delete_range(doc, cn_resp_start + 1, cn_proj_start - 1)

cn_resp_start = find_heading(doc, '职位及成就')

# ---- Build Chinese Responsibilities content ----
cn_resp_paras = []

# Thoughtworks
cn_resp_paras.append(make_para(doc, '2021 – 2025           公司名称：Thoughtworks', font=FONT))
cn_resp_paras.append(make_para(doc, '职位名称：前端开发工程师', font=FONT))
cn_resp_paras.append(make_para(doc, ''))
cn_resp_paras.append(make_para(doc, '主要职责:', bold=True, font=FONT))

cn_resp_paras.append(make_list_item(doc, 'REA Group 全栈开发与AI工作流：负责前端功能开发及Node.js/Express后端服务维护；在Codex中设计并实现了4个自定义Agent Skill（Spec需求拆分 / Unit Test单测生成 / Issue报告发布 / Code Review预检），覆盖从需求拆解→开发→单测→Issue发布→Code Review→自主触发CI/CD发布的完整交付链路。', style='List Paragraph'))
cn_resp_paras.append(make_list_item(doc, '保时捷多端经销商系统：负责微信小程序、抖音小程序、Web Admin、H5四端核心功能开发；主导车辆智能推荐模块、新车模块、车系模块的设计实现；负责Admin端经销商管理（列表、审核、权限、数据看板）及腾讯云视频接入；主导微信→抖音小程序迁移，解决Taro API兼容及动画性能问题，实现约70%代码复用；参与跨端组件库建设（Taro + Lerna Monorepo），封装15+通用业务组件。', style='List Paragraph'))
cn_resp_paras.append(make_list_item(doc, '推行ESLint规范及Code Review机制，团队生产缺陷率降低约50%。', style='List Paragraph'))
cn_resp_paras.append(make_list_item(doc, '租赁系统：开发Web端订单管理（列表、筛选、详情）及小程序选车功能；编写Jest + React Testing Library单元及集成测试。', style='List Paragraph'))
cn_resp_paras.append(make_list_item(doc, '内部系统：开发KPI可视化看板（ECharts），替代人工Excel报表；开发合同管理系统（列表、筛选、批量操作）；重构遗留代码，包体积减少约35%。', style='List Paragraph'))
cn_resp_paras.append(make_para(doc, ''))

cn_resp_paras.append(make_para(doc, '2020 – 2021           公司名称：西安飞蜂科技', font=FONT))
cn_resp_paras.append(make_para(doc, '职位名称：前端开发工程师', font=FONT))
cn_resp_paras.append(make_para(doc, ''))
cn_resp_paras.append(make_para(doc, '主要职责:', bold=True, font=FONT))
cn_resp_paras.append(make_list_item(doc, '独立开发乡村振兴公厕管理小程序（原生微信小程序）：集成IoT传感器数据实时监测、动态路径规划、车辆GPS实时追踪与导航。', style='List Paragraph'))
cn_resp_paras.append(make_list_item(doc, '独立完成内部管理系统全栈开发（React + Spring Boot + MongoDB），从数据库建模到前端UI。', style='List Paragraph'))

# Insert
insert_paras_after(doc.paragraphs[cn_resp_start], cn_resp_paras)
print(f"Inserted {len(cn_resp_paras)} paragraphs for Chinese 职位及成就")

# ============================================================
# STEP 4: Chinese "项目经验"
# ============================================================
cn_proj_start = find_heading(doc, '项目经验')
# Find end of document or next section
# Check if there's content after 项目经验
total = len(doc.paragraphs)
delete_range(doc, cn_proj_start + 1, total - 1)

cn_proj_start = find_heading(doc, '项目经验')

# ---- Build Chinese Project Experience ----
cn_proj_paras = []

# Project 1
cn_proj_paras.append(make_para(doc, '2022 – 2025       保时捷多端经销商系统', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '项目描述：为保时捷全国经销商网络开发的核心业务系统，覆盖微信小程序、抖音小程序、Web Admin、H5四端。包含车辆智能推荐、新车目录、经销商管理、培训视频等功能模块。', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '主要职责:', bold=True, font=FONT))
cn_proj_paras.append(make_list_item(doc, '四端核心业务模块架构设计与开发，确保跨平台一致的UX和共享业务逻辑', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '主导微信→抖音小程序迁移：解决Taro API不兼容问题，重写动画模块适配抖音运行时，实现约70%代码复用', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '参与Taro + Lerna Monorepo跨端组件库建设，封装15+通用业务组件，消除多端重复开发', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, 'Admin端经销商管理：列表、审核流程、权限分配、数据看板；集成腾讯云视频', style='List Paragraph'))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '项目成果：', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '通过组件复用，多端交付周期缩短约40%。跨平台UI一致性达到100%。', font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '核心技术：React, TypeScript, Redux Toolkit, TaroJS, 微信小程序, 抖音小程序, H5, 腾讯云视频', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, ''))

# Project 2
cn_proj_paras.append(make_para(doc, '2023 – 2025       REA Group — AI驱动全栈开发', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '项目描述：维护和开发澳大利亚最大地产平台（realcommercial.com.au / realestate.com.au），服务数百万用户。全英文与澳洲、印度、中国团队协作。设计并实现自定义Agent Skill驱动的AI开发流水线。', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '主要职责:', bold=True, font=FONT))
cn_proj_paras.append(make_list_item(doc, '前端功能开发（React + TypeScript）及Node.js/Express后端服务维护', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '设计4个自定义Agent Skill：Spec（需求拆分）、Unit Test（Jest单测自动生成）、Issue（结构化根因报告）、Code Review（自动化预检）', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '完整交付管线：需求拆解 → 开发 → 单测 → Issue发布 → Code Review → 自主触发CI/CD发布', style='List Paragraph'))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '项目成果：', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '将手工、碎片化的开发流程转变为可复制、有质量门禁的AI辅助交付体系。每次变更包含强制测试覆盖、根因文档和正式Review。', font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '核心技术：React, TypeScript, Node.js, Express, Jest, Codex (AI辅助)', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, ''))

# Project 3
cn_proj_paras.append(make_para(doc, '2020 – 2021       乡村振兴公厕IoT管理小程序', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '项目描述：乡村振兴项目——在农村公厕安装IoT传感器，实时监测厕位状态。小程序帮助政府工作人员查看需要清理的设施，并按动态规划的路线进行作业。', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '主要职责:', bold=True, font=FONT))
cn_proj_paras.append(make_list_item(doc, '使用微信原生框架独立开发完整小程序', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '集成IoT传感器数据，实时展示设施状态监测', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '实现动态路径规划算法，根据传感器优先级和地理位置生成最优收集路径', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '构建车辆GPS实时追踪及导航功能', style='List Paragraph'))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '项目成果：', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '实现数据驱动的环卫管理，替代人工排班模式。', font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '核心技术：原生微信小程序, IoT传感器集成, GPS追踪, 路径规划', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, ''))

# Project 4
cn_proj_paras.append(make_para(doc, '2020 – 2021       全栈内部管理系统', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '项目描述：独立设计并交付完整的内部管理系统，覆盖数据库设计到前端UI的全栈开发。', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '主要职责:', bold=True, font=FONT))
cn_proj_paras.append(make_list_item(doc, '设计MongoDB数据模型，使用Spring Boot构建RESTful API', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '开发React前端，实现完整CRUD功能', style='List Paragraph'))
cn_proj_paras.append(make_list_item(doc, '负责部署及生产上线', style='List Paragraph'))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '项目成果：', bold=True, font=FONT))
cn_proj_paras.append(make_para(doc, '独立完成全栈系统从零到上线的完整交付。', font=FONT))
cn_proj_paras.append(make_para(doc, ''))
cn_proj_paras.append(make_para(doc, '核心技术：React, Spring Boot, MongoDB', bold=True, font=FONT))

# Insert
insert_paras_after(doc.paragraphs[cn_proj_start], cn_proj_paras)
print(f"Inserted {len(cn_proj_paras)} paragraphs for Chinese 项目经验")

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/panlinfeng/workspace/course-1-basic-stuff/简历/Linfeng Pan.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
